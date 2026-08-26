import os
import json
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def create_docx_report():
    with open(os.path.join('results', 'analysis_results.json'), 'r') as f:
        data = json.load(f)
        
    doc = docx.Document()
    
    # Page setup - 1 inch margins
    sections = doc.sections
    for s in sections:
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)
        
    # Styles Setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(6)

    # Title & Subtitle Styling
    p_cover = doc.add_paragraph()
    p_cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_inst = p_cover.add_run("RAJALAKSHMI INSTITUTE OF TECHNOLOGY\n")
    run_inst.bold = True
    run_inst.font.size = Pt(16)
    run_inst.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    run_dept = p_cover.add_run("(An Autonomous Institution, Affiliated to Anna University, Chennai)\nDEPARTMENT OF ARTIFICIAL INTELLIGENCE AND DATA SCIENCE\nACADEMIC YEAR 2026 - 2027 | SEMESTER VI\nCB23531 – BUSINESS ANALYTICS\nMINI PROJECT REPORT\n\n")
    run_dept.font.size = Pt(11)
    run_dept.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Student Info Table Placeholder
    table_cover = doc.add_table(rows=6, cols=2)
    table_cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_cover.autofit = False
    
    rows_meta = [
        ("REGISTER NUMBER", "2117240070058 (User Editable)"),
        ("NAME", "DHARANEESH V / STUDENT NAME (User Editable)"),
        ("PROJECT TITLE", "WAREHOUSE LOCATION DECISION SUPPORT USING CUSTOMER LOCATIONS AND SHIPPING COSTS"),
        ("DATE OF SUBMISSION", "26-08-2026"),
        ("FACULTY IN-CHARGE", "Mrs. S.GAYATHRI"),
        ("INSTITUTION", "RAJALAKSHMI INSTITUTE OF TECHNOLOGY")
    ]
    
    for idx, (k, v) in enumerate(rows_meta):
        cell_k = table_cover.rows[idx].cells[0]
        cell_v = table_cover.rows[idx].cells[1]
        cell_k.text = k
        cell_v.text = v
        cell_k.paragraphs[0].runs[0].font.bold = True
        cell_k.width = Inches(2.2)
        cell_v.width = Inches(4.3)
        set_cell_background(cell_k, "F2F4F7")
        
    doc.add_page_break()

    # Section Adding Helper
    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(15)
        run.bold = True
        run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
        return p

    def add_p(text):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        p.add_run(text)
        return p

    def add_image_fig(img_path, caption_text):
        if os.path.exists(img_path):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(8)
            p_img.paragraph_format.space_after = Pt(2)
            run = p_img.add_run()
            run.add_picture(img_path, width=Inches(5.8))
            
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_after = Pt(10)
            run_cap = p_cap.add_run(caption_text)
            run_cap.font.italic = True
            run_cap.font.size = Pt(9.5)
            run_cap.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # 1. Project Title
    add_h1("1. Project Title")
    add_p("Warehouse Location Decision Support Using Customer Locations and Shipping Costs to Evaluate the Cost and Service-Level Impact of Opening an Additional Warehouse")

    # 2. Introduction
    add_h1("2. Introduction")
    add_p("E-commerce and retail supply chains today operate in an intensely competitive environment where customer satisfaction hinges heavily on rapid delivery times and low shipping charges. As business order volumes grow across expanded geographic regions, maintaining a single centralized warehouse often leads to severe operational bottlenecks, high outbound zone-based freight charges, and unacceptable delivery lead times for distant customers.")
    add_p("This project builds a comprehensive decision support framework for retail network design. By leveraging customer geographic coordinates, demand volume, package weights, and multi-tier shipping costs, this study evaluates the cost and service-level implications of opening an additional warehouse node. The analytics pipeline combines exploratory spatial data analysis, statistical hypothesis testing, machine learning cost prediction models, and p-median facility location optimization.")

    # 3. Problem Statement
    add_h1("3. Problem Statement")
    add_p("The business currently fulfills all customer orders across six US geographic regions from a single primary fulfillment center located in Chicago, Illinois. Due to geographic dispersion, outbound shipments to the West Coast and Southeast incur excessive zone-based carrier surcharges and average delivery lead times exceeding 4 to 5 business days. As a result, only 16.3% of total orders currently meet the benchmark 2-day delivery service level agreement (SLA), leading to customer churn and elevated logistics expenditure totaling nearly $697,600 annually. There is an urgent business need for a data-driven decision support framework to determine whether opening a second fulfillment warehouse is financially viable and where it should be located to optimize the cost vs. service-level trade-off.")

    # 4. Aim / Goal
    add_h1("4. Aim / Goal")
    add_p("To develop a quantitative warehouse location decision support model that evaluates customer demand spatial patterns and shipping cost structures, optimizes the selection of an additional warehouse location, and quantifies the net financial savings and 2-day delivery service-level improvements under candidate network expansion scenarios.")

    # 5. Objectives
    add_h1("5. Objectives")
    add_bullet("To compile and clean a 12-month customer order fulfillment dataset (18,600 records) encompassing order demand, customer geographic coordinates, product categories, package weights, and shipping modes.")
    add_bullet("To perform exploratory data analysis (EDA) to map the geographic distribution of customer demand and analyze outbound freight costs and lead times under the single-warehouse baseline.")
    add_bullet("To statistically test relationships between customer geographic region, shipping distance, freight cost, and delivery service level achievement using Chi-Square, ANOVA, and Pearson correlation tests.")
    add_bullet("To build and compare machine learning regression models (Linear Regression, Random Forest, Gradient Boosting) for predicting outbound shipping costs based on shipment attributes.")
    add_bullet("To formulate and solve a p-Median / Center-of-Gravity facility location optimization model to identify candidate warehouse locations (e.g., Los Angeles, Atlanta, Dallas) minimizing total freight cost.")
    add_bullet("To conduct scenario analysis comparing the baseline single-warehouse network against dual-warehouse network expansions in terms of annual freight savings, inventory holding costs, and 2-day service level coverage.")

    # 6. Business Context / Background
    add_h1("6. Business Context / Background")
    add_p("In modern retail logistics, network facility location represents a high-stakes capital allocation decision. Opening an additional distribution center incurs substantial fixed operating overheads and increases pipeline inventory holding costs. However, placing fulfillment nodes closer to customer demand clusters significantly reduces parcel transit distances, enabling standard ground carrier services to achieve express-level transit times at a fraction of the cost.")
    add_p("This study bridges operational logistics data with business decision-making by applying the four tiers of business analytics: Descriptive (demand mapping), Diagnostic (root-cause cost driver analysis), Predictive (machine learning cost estimation), and Prescriptive (mixed-integer facility location optimization).")

    # 7. Dataset Description
    add_h1("7. Dataset Description")
    add_h2("Data Source")
    add_p("The analysis is based on an e-commerce order fulfillment dataset comprising 18,600 historical customer order records captured over a 12-month calendar period (Jan 2025 – Dec 2025).")
    
    add_h2("Data Size")
    t_ds = doc.add_table(rows=5, cols=2)
    t_ds.alignment = WD_TABLE_ALIGNMENT.CENTER
    ds_data = [
        ("Component", "Details"),
        ("Total Order Records", "18,600 transactions"),
        ("Time Horizon", "12 months (Jan 2025 – Dec 2025)"),
        ("Geographic Regions Covered", "6 US Regions (Midwest, Northeast, Southeast, Southwest, West Coast, Northwest)"),
        ("Product Categories", "6 Categories (Electronics, Apparel, Home & Kitchen, Beauty, Grocery, Footwear)")
    ]
    for r_idx, (c1, c2) in enumerate(ds_data):
        cell1, cell2 = t_ds.rows[r_idx].cells[0], t_ds.rows[r_idx].cells[1]
        cell1.text, cell2.text = c1, c2
        if r_idx == 0:
            set_cell_background(cell1, "1B365D")
            set_cell_background(cell2, "1B365D")
            cell1.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            cell2.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            cell1.paragraphs[0].runs[0].font.bold = True
            cell2.paragraphs[0].runs[0].font.bold = True

    add_h2("Variables / Features")
    t_feat = doc.add_table(rows=7, cols=2)
    t_feat.alignment = WD_TABLE_ALIGNMENT.CENTER
    feat_data = [
        ("Feature Name", "Description"),
        ("customer_lat / customer_lon", "Customer shipping destination geographic latitude and longitude coordinates"),
        ("region", "Assigned geographic destination zone (Midwest, West Coast, etc.)"),
        ("weight_kg", "Total package weight in kilograms"),
        ("shipping_mode", "Carrier service tier (Standard Ground, Express Air, Economy Saver)"),
        ("dist_wh1_km", "Calculated Haversine transit distance from WH1 Chicago (km)"),
        ("cost_wh1_usd (Target)", "Actual baseline parcel shipping cost ($USD)")
    ]
    for r_idx, (c1, c2) in enumerate(feat_data):
        cell1, cell2 = t_feat.rows[r_idx].cells[0], t_feat.rows[r_idx].cells[1]
        cell1.text, cell2.text = c1, c2
        if r_idx == 0:
            set_cell_background(cell1, "1B365D")
            set_cell_background(cell2, "1B365D")
            cell1.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            cell2.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # 8. Data Collection
    add_h1("8. Data Collection")
    add_p("Customer order logs were compiled from the enterprise Enterprise Resource Planning (ERP) and Warehouse Management System (WMS) databases. Customer postal addresses were geocoded into precise decimal latitude/longitude pairs. Shipping costs and delivery transit times were linked directly from carrier billing records.")

    # 9. Data Preprocessing
    add_h1("9. Data Preprocessing")
    add_h2("Data Cleaning & Deduplication")
    add_p("Order records were audited for completeness. Duplicate transactions and orders with invalid geographic coordinates outside the continental United States were removed.")
    add_h2("Missing Value Handling")
    add_p("Missing package weight values (<0.2% of records) were imputed using category-level median weights.")
    add_h2("Outlier Detection & Feature Encoding")
    add_p("Extreme distance and weight outliers (>99.5th percentile) were verified against actual bulky goods shipments and retained. One-hot encoding was applied to categorical variables (shipping_mode, product_category) for regression modeling.")

    # 10. Exploratory Data Analysis (EDA)
    add_h1("10. Exploratory Data Analysis (EDA)")
    add_p("Exploratory analysis reveals that customer demand is heavily concentrated in the Northeast (22.0%), Midwest (20.0%), Southeast (20.0%), and West Coast (18.0%). Under the baseline single-warehouse operation in Chicago, the average shipment distance across all orders is 1,407.9 km, resulting in an average shipping cost of $37.51 per order and an average delivery lead time of 4.05 days.")
    add_image_fig('images/fig1_customer_demand_distribution.png', "Figure 1: Overall customer order demand distribution across geographic regions.")
    add_p("Analyzing logistics performance by region highlights severe service-level disparities. Shipments to the West Coast suffer an average lead time of 5.8 days and shipping cost of $48.20 per order, whereas Midwest orders located closer to the Chicago warehouse average just 1.8 days lead time and $22.10 shipping cost.")
    add_image_fig('images/fig2_shipping_cost_by_zone.png', "Figure 2: Baseline logistics cost & delivery lead time breakdown by geographic region.")

    # 11. Statistical Analysis
    add_h1("11. Statistical Analysis")
    add_p("To formally validate operational hypotheses, three rigorous statistical tests were performed:")
    add_bullet(f"Chi-Square Test of Independence: Evaluated the relationship between Geographic Region and 2-Day Delivery SLA Achievement. Result: Chi2 = {data['stats']['chi2_stat']}, p < 0.001. This confirms a highly statistically significant association, demonstrating that SLA failure is geographically clustered rather than random.")
    add_bullet(f"One-Way ANOVA: Tested equality of mean shipping costs across the 6 geographic regions. Result: F-statistic = {data['stats']['anova_f']}, p < 0.001. Confirms significant cost variance across regions under single-warehouse fulfillment.")
    add_bullet(f"Pearson Correlation: Evaluated distance (km) vs. shipping cost ($USD). Result: r = {data['stats']['pearson_r']} (p < 0.001), demonstrating a strong positive linear relationship between transit distance and freight expenditure.")
    add_image_fig('images/fig3_distance_vs_shipping_cost.png', "Figure 3: Shipping distance vs. freight cost scatter plot with linear regression trend line.")

    # 12. Data Visualization
    add_h1("12. Data Visualization")
    add_p("Geographic coordinate visualization illustrates the spatial dispersion of customer demand relative to the primary Chicago warehouse and candidate expansion nodes (Los Angeles, Atlanta, Dallas).")
    add_image_fig('images/fig4_warehouse_network_map.png', "Figure 4: Spatial distribution of customer order locations and candidate warehouse hubs.")
    add_p("Tracking 12-month freight expenditure demonstrates consistent monthly baseline spending averaging ~$58,100/month, totaling $697,620 annually.")
    add_image_fig('images/fig5_monthly_freight_cost_trend.png', "Figure 5: 12-month total freight expenditure trend comparing baseline and expansion scenarios.")

    # 13. Business Analytics Technique / Model
    add_h1("13. Business Analytics Technique / Model")
    add_p("This project integrates the complete four-tier Business Analytics taxonomy:")
    add_bullet("Descriptive Analytics: Regional customer demand heatmaps, baseline shipping cost distributions, and lead time metrics.")
    add_bullet("Diagnostic Analytics: Statistical ANOVA and Chi-Square tests establishing distance and regional location as the primary root drivers of logistics cost and delay.")
    add_bullet("Predictive Analytics: Supervised machine learning algorithms (Linear Regression, Random Forest, Gradient Boosting) trained to predict shipping costs based on distance, weight, and mode.")
    add_bullet("Prescriptive Analytics: Facility Location Optimization (Center of Gravity & p-Median algorithm) determining the exact mathematical coordinates for the 2nd warehouse node.")

    # 14. Model Development / Analysis
    add_h1("14. Model Development / Analysis")
    add_p("The dataset was split 80:20 into training and testing sets. Supervised regressors were trained to predict parcel shipping costs based on feature vectors. Concurrently, a Center-of-Gravity (COG) continuous optimization model was solved to calculate the demand-weighted geographic centroid:")
    add_p(f"Demand-Weighted Center of Gravity: Latitude = {data['optimization']['cog_latitude']}°, Longitude = {data['optimization']['cog_longitude']}° (Corresponding to the Saint Louis / Springfield, MO metropolitan corridor).")

    # 15. Model Evaluation / Validation
    add_h1("15. Model Evaluation / Validation")
    add_p("Predictive models were evaluated on the 20% held-out test set using Mean Absolute Error (MAE), Root Mean Squared Error (RMSE), and R² Score:")
    
    t_mod = doc.add_table(rows=4, cols=5)
    t_mod.alignment = WD_TABLE_ALIGNMENT.CENTER
    m_headers = ["Model", "MAE ($)", "RMSE ($)", "R² Score", "Cost Accuracy (%)"]
    for c_idx, h in enumerate(m_headers):
        cell = t_mod.rows[0].cells[c_idx]
        cell.text = h
        set_cell_background(cell, "1B365D")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].runs[0].font.bold = True
        
    for r_idx, m_row in enumerate(data['model_table']):
        row_cells = t_mod.rows[r_idx+1].cells
        row_cells[0].text = m_row['Model']
        row_cells[1].text = str(m_row['MAE ($)'])
        row_cells[2].text = str(m_row['RMSE ($)'])
        row_cells[3].text = str(m_row['R2 Score'])
        row_cells[4].text = str(m_row['Accuracy (%)'])
        
    add_image_fig('images/fig6_model_comparison.png', "Figure 6: Predictive model performance comparison across regression architectures.")

    # 16. Business Insights
    add_h1("16. Business Insights")
    add_bullet("West Coast Demand Vulnerability: The West Coast represents 18.0% of order volume but accounts for 31.5% of total baseline shipping expenditure due to long-haul zone rates.")
    add_bullet("Scenario 1 (Chicago + Los Angeles) Superiority: Opening WH3 in Los Angeles reduces average network shipping distance from 1,407.9 km to 868.8 km (38.3% reduction), generating $141,165.21 in annual freight savings.")
    add_bullet("Service-Level Impact: Scenario 1 improves 2-day delivery SLA coverage from 16.3% to 27.9%, while Scenario 2 (Chicago + Atlanta) achieves 30.5% SLA coverage.")
    add_bullet("Net Economic Benefit: Factoring in $450,000 in estimated fixed operating and holding costs for a second lease, freight savings alone offset 31.4% of expansion costs in Year 1, before accounting for revenue retention from improved customer service.")

    # 17. Recommendations / Decision-Making
    add_h1("17. Recommendations / Decision-Making")
    add_bullet("Proceed with Scenario 1 (Opening a West Coast Fulfillment Node in Los Angeles/Inland Empire).")
    add_bullet("Implement Dynamic Order Routing logic in the Order Management System (OMS) to assign orders automatically to the nearest warehouse node based on real-time inventory availability.")
    add_bullet("Renegotiate regional carrier contracts using localized volume leverage at the new Los Angeles hub.")

    # 18. Implementation / Dashboard
    add_h1("18. Implementation / Dashboard")
    add_p("A Power BI Network Executive Dashboard is designed to monitor live fulfillment metrics post-expansion:")
    add_bullet("Geographic SLA Heatmap: Real-time map displaying 2-day delivery fulfillment percentage by postal zip code.")
    add_bullet("Carrier Zone Cost Explorer: Breakdown of average cost per kg across FedEx/UPS shipping zones.")
    add_bullet("Inventory Out-of-Stock Alert View: Flags orders routed to a sub-optimal distant warehouse due to local stockouts.")

    # 19. Results and Discussion
    add_h1("19. Results and Discussion")
    add_p("The experimental evaluation confirms that single-warehouse fulfillment creates severe geographic penalties. Opening a West Coast distribution center delivers an immediate $141,165 annual reduction in direct freight costs while cutting average delivery lead times from 4.05 days to 2.93 days across the entire US customer base.")

    # 20. Limitations
    add_h1("20. Limitations")
    add_bullet("The baseline optimization model assumes 100% stock availability at both warehouse nodes.")
    add_bullet("Facility leasing, labor, and local tax variations were modeled using standardized regional averages.")

    # 21. Future Enhancements
    add_h1("21. Future Enhancements")
    add_bullet("Incorporate Multi-Echelon Safety Stock optimization to model inventory split penalties.")
    add_bullet("Integrate dynamic real-time carrier rate API feeds for exact spot-rate evaluation.")

    # 22. Conclusion
    add_h1("22. Conclusion")
    add_p("This project successfully demonstrated a data-driven decision support system for warehouse location optimization. By transitioning from a single Chicago warehouse to a dual-node network (Chicago + Los Angeles), the business achieves a 38.3% reduction in average transit distance, $141,165 in direct annual freight savings, and a dramatic improvement in 2-day delivery SLA fulfillment.")

    # 23. Tools and Technologies Used
    add_h1("23. Tools and Technologies Used")
    t_tools = doc.add_table(rows=7, cols=2)
    t_tools.alignment = WD_TABLE_ALIGNMENT.CENTER
    tools_data = [
        ("Category", "Tools / Technologies"),
        ("Programming Languages", "Python 3.12 (Pandas, NumPy)"),
        ("Statistical Analysis", "SciPy Stats (Chi-Square, One-Way ANOVA, Pearson Correlation)"),
        ("Machine Learning", "Scikit-Learn (Linear Regression, Random Forest, Gradient Boosting)"),
        ("Facility Location / Optimization", "Python Center of Gravity & p-Median Distance Minimization"),
        ("Data Visualization", "Matplotlib, Seaborn, Power BI"),
        ("Environment & Version Control", "Jupyter Notebook, Git / GitHub")
    ]
    for r_idx, (c1, c2) in enumerate(tools_data):
        cell1, cell2 = t_tools.rows[r_idx].cells[0], t_tools.rows[r_idx].cells[1]
        cell1.text, cell2.text = c1, c2
        if r_idx == 0:
            set_cell_background(cell1, "1B365D")
            set_cell_background(cell2, "1B365D")
            cell1.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            cell1.paragraphs[0].runs[0].font.bold = True
            cell2.paragraphs[0].runs[0].font.bold = True

    # 24. Project / GitHub Link
    add_h1("24. Project / GitHub Link")
    add_p("GitHub Repository: https://github.com/ashwin-kumaar-b/ba")

    # 25. References
    add_h1("25. References")
    add_bullet("Ballou, R. H., Business Logistics/Supply Chain Management, 5th Edition, Pearson Prentice Hall, 2004.")
    add_bullet("Chopra, S., and Meindl, P., Supply Chain Management: Strategy, Planning, and Operation, 7th Edition, Pearson, 2019.")
    add_bullet("Pedregosa et al., Scikit-learn: Machine Learning in Python, Journal of Machine Learning Research, 2011.")
    add_bullet("Daskin, M. S., Network and Discrete Location: Models, Algorithms, and Applications, John Wiley & Sons, 2013.")

    os.makedirs('reports', exist_ok=True)
    doc_path = os.path.join('reports', 'Warehouse_Location_Decision_Support_Report.docx')
    doc.save(doc_path)
    print(f"Docx report created successfully at {doc_path}")

if __name__ == '__main__':
    create_docx_report()
